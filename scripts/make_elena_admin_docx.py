#!/usr/bin/env python3
"""Собрать ELENA-ADMIN.docx из инструкции для Елены."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ELENA-ADMIN.docx"


def set_default_font(doc: Document, name: str = "Calibri", size: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x3D, 0x2E)


def add_lead(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.bold = True


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3D, 0x2E)


def add_steps(doc: Document, steps: list[str]) -> None:
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_menu_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Раздел"
    hdr[1].text = "Что там"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for i, (section, desc) in enumerate(rows, 1):
        row = table.rows[i].cells
        row[0].text = section
        row[1].text = desc


def build() -> None:
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    add_title(doc, "Как пользоваться админкой сайта")
    add_lead(doc, "Одна закладка в браузере:")
    add_lead(doc, "https://ваш-домен.ru/admin/")
    add_note(doc, "(замените на реальный адрес после запуска сайта)")
    doc.add_paragraph()

    add_heading(doc, "Вход")
    add_steps(
        doc,
        [
            "Откройте закладку «Админка сайта» в Chrome или Яндекс.Браузере.",
            "Нажмите «Войти» / «Log in».",
            "Введите ваш email (тот, на который вас пригласили).",
            "На почту придёт письмо — откройте его и нажмите ссылку для входа. Пароль придумывать не нужно.",
        ],
    )
    add_note(doc, "Если письма нет — проверьте папку «Спам».")
    doc.add_paragraph()

    add_heading(doc, "Главное меню слева")
    add_menu_table(
        doc,
        [
            ("Экскурсии", "Маршруты на странице экскурсий"),
            ("Подкаст", "Выпуски подкаста"),
            ("Книги", "Книги и ссылки на OZON"),
            ("Журнал", "Статьи и новости"),
            ("Контакты", "Email, телефон, Telegram, ссылки VK"),
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "Добавить новый выпуск подкаста")
    add_steps(
        doc,
        [
            "Слева: Подкаст → Выпуски.",
            "Прокрутите вниз, нажмите «Добавить выпуск» (или «+»).",
            "Заполните: название, дату (например 11.06.2026), краткое описание, обложку (кнопка «Выбрать»), ссылку на страницу выпуска со страницы MediaMetrics.",
            "Поле «Ссылка для встроенного видео» — если не знаете, оставьте пустым и напишите разработчику: он вставит сам.",
            "Убедитесь, что «Показывать на сайте» включено.",
            "Нажмите «Опубликовать» вверху справа.",
            "Подождите 2–3 минуты и обновите страницу подкаста на сайте (F5).",
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "Добавить экскурсию")
    add_steps(
        doc,
        [
            "Экскурсии → Маршруты → Добавить экскурсию.",
            "Заполните название, описание, фото, длительность, цену.",
            "Нажмите «Опубликовать».",
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "Добавить статью в журнал")
    add_steps(
        doc,
        [
            "Журнал → Статьи и новости → Добавить материал.",
            "Заполните заголовок, краткое описание, основной текст (абзацы через пустую строку).",
            "Нажмите «Опубликовать».",
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "Скрыть материал, не удаляя")
    p = doc.add_paragraph()
    p.add_run("Снимите галочку «Показывать на сайте» → «Опубликовать». ")
    p.add_run("На сайте карточка исчезнет, данные сохранятся.")
    doc.add_paragraph()

    add_heading(doc, "Если что-то пошло не так")
    add_bullets(
        doc,
        [
            "Не получается войти → напишите разработчику, он пришлёт новое приглашение.",
            "Нажали не то → не закрывайте вкладку, напишите в Telegram.",
            "Сайт не обновился через 5 минут → обновите страницу (F5) или напишите разработчику.",
        ],
    )
    add_note(
        doc,
        "Не трогайте в разделе «Контакты» блоки «SEO и домен» и «Аналитика» — это для разработчика.",
    )
    doc.add_paragraph()

    add_heading(doc, "Что НЕ нужно делать")
    add_bullets(
        doc,
        [
            "Не открывать папки на компьютере и не искать файлы .json",
            "Не пользоваться FTP и reg.ru",
            "Не удалять поля «Служебное имя (slug)» у старых записей — можно оставить как есть",
        ],
    )
    p = doc.add_paragraph()
    run = p.add_run("Всё редактирование — только через эту админку в браузере.")
    run.bold = True

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
